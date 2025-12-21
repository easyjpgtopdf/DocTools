"""
FREE PDF to Word Pipeline - Data-Oriented Layout Reconstruction
Improves text-based PDF conversion without visual features (no images, borders, colors).

Features:
- X-axis column detection
- Improved row detection (y-distance + font-size similarity)
- Multi-line cell merge
- Page-wise table isolation
- Header/footer suppression
- Visual PDF detection (blocks conversion if detected)
"""
import logging
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass
from collections import defaultdict
from docx import Document
from docx.shared import Pt
from pdfminer.layout import LAParams, LTTextBox, LTTextLine, LTChar, LTFigure
from pdfminer.high_level import extract_pages

logger = logging.getLogger(__name__)


@dataclass
class TextItem:
    """Text item with position and metadata."""
    text: str
    x0: float  # Left
    y0: float  # Bottom
    x1: float  # Right
    y1: float  # Top
    font_size: float
    page_num: int


@dataclass
class Column:
    """Column definition."""
    x0: float
    x1: float
    index: int  # 0-based index, left to right


@dataclass
class Cell:
    """Table cell."""
    text: str
    col_index: int
    row_index: int
    x0: float
    x1: float
    y0: float
    y1: float
    font_size: float


def extract_text_with_positions(pdf_path: str) -> List[TextItem]:
    """
    Extract text items with positions from PDF using pdfminer.
    
    Args:
        pdf_path: Path to PDF file
        
    Returns:
        List of TextItem objects with positions
    """
    text_items = []
    
    try:
        laparams = LAParams(
            line_margin=0.5,
            word_margin=0.1,
            char_margin=2.0,
            boxes_flow=0.5
        )
        
        for page_num, page_layout in enumerate(extract_pages(pdf_path, laparams=laparams), start=1):
            page_height = page_layout.height
            page_width = page_layout.width
            
            for element in page_layout:
                if isinstance(element, LTTextBox):
                    for line in element:
                        if isinstance(line, LTTextLine):
                            # Get bounding box
                            x0, y0, x1, y1 = line.bbox
                            
                            # Get text
                            text = line.get_text().strip()
                            if not text:
                                continue
                            
                            # Get font size from first character
                            font_size = 12.0  # Default
                            for char in line:
                                if isinstance(char, LTChar):
                                    try:
                                        # Try to get font size from character
                                        if hasattr(char, 'height'):
                                            font_size = char.height
                                        elif hasattr(char, 'size'):
                                            font_size = char.size
                                        break
                                    except:
                                        pass
                            
                            # Convert Y coordinates (pdfminer uses bottom-left, we need top-left)
                            y0_converted = page_height - y1
                            y1_converted = page_height - y0
                            
                            text_items.append(TextItem(
                                text=text,
                                x0=x0,
                                y0=y0_converted,
                                x1=x1,
                                y1=y1_converted,
                                font_size=font_size,
                                page_num=page_num
                            ))
    
    except Exception as e:
        logger.error(f"Error extracting text with positions: {e}", exc_info=True)
        raise
    
    return text_items


def detect_columns(text_items: List[TextItem], page_num: int, tolerance: float = 10.0) -> List[Column]:
    """
    Detect columns by clustering X positions.
    
    Args:
        text_items: List of text items (filtered to one page)
        page_num: Page number
        tolerance: X-position tolerance for column clustering
        
    Returns:
        List of Column objects sorted left to right
    """
    # Filter items for this page
    page_items = [item for item in text_items if item.page_num == page_num]
    
    if not page_items:
        return []
    
    # Extract all X positions (left edges)
    x_positions = sorted(set([item.x0 for item in page_items]))
    
    # Cluster X positions into columns
    columns = []
    current_cluster = []
    
    for x in x_positions:
        if not current_cluster or abs(x - current_cluster[-1]) <= tolerance:
            current_cluster.append(x)
        else:
            # Save previous cluster
            if current_cluster:
                avg_x = sum(current_cluster) / len(current_cluster)
                columns.append(avg_x)
            current_cluster = [x]
    
    if current_cluster:
        avg_x = sum(current_cluster) / len(current_cluster)
        columns.append(avg_x)
    
    # Create Column objects sorted left to right
    sorted_columns = sorted(columns)
    column_objects = []
    
    for idx, x in enumerate(sorted_columns):
        # Find right edge (next column's left edge or page width)
        if idx + 1 < len(sorted_columns):
            x1 = sorted_columns[idx + 1] - tolerance
        else:
            # Last column extends to max x1
            x1 = max(item.x1 for item in page_items)
        
        column_objects.append(Column(
            x0=x,
            x1=x1,
            index=idx
        ))
    
    return column_objects


def assign_column(text_item: TextItem, columns: List[Column]) -> Optional[int]:
    """
    Assign text item to a column.
    
    Args:
        text_item: Text item
        columns: List of columns
        
    Returns:
        Column index or None if no match
    """
    if not columns:
        return None
    
    # Find column where item's x0 falls
    for col in columns:
        if col.x0 <= text_item.x0 <= col.x1:
            return col.index
        # Also check if item overlaps with column
        if text_item.x0 < col.x1 and text_item.x1 > col.x0:
            return col.index
    
    # If no match, assign to nearest column
    distances = [abs(text_item.x0 - col.x0) for col in columns]
    min_idx = distances.index(min(distances))
    return columns[min_idx].index


def detect_rows(text_items: List[TextItem], page_num: int, font_size_tolerance: float = 2.0, y_tolerance_factor: float = 0.5) -> List[Dict]:
    """
    Detect rows using y-distance + font-size similarity.
    
    Args:
        text_items: List of text items (filtered to one page)
        page_num: Page number
        font_size_tolerance: Font size difference tolerance
        y_tolerance_factor: Y-distance tolerance as factor of font size
        
    Returns:
        List of row dictionaries with items and metadata
    """
    # Filter items for this page
    page_items = [item for item in text_items if item.page_num == page_num]
    
    if not page_items:
        return []
    
    # Sort by Y position (top to bottom)
    page_items.sort(key=lambda x: x.y0, reverse=True)
    
    rows = []
    current_row = None
    
    for item in page_items:
        if current_row is None:
            # Start new row
            current_row = {
                'items': [item],
                'y_center': (item.y0 + item.y1) / 2,
                'font_size': item.font_size,
                'y_min': item.y0,
                'y_max': item.y1
            }
        else:
            # Check if item belongs to current row
            item_y_center = (item.y0 + item.y1) / 2
            y_distance = abs(item_y_center - current_row['y_center'])
            font_diff = abs(item.font_size - current_row['font_size'])
            
            # Dynamic Y tolerance based on font size
            avg_font_size = (item.font_size + current_row['font_size']) / 2
            y_tolerance = avg_font_size * y_tolerance_factor
            
            # Item belongs to current row if:
            # 1. Y distance is small (similar row)
            # 2. Font size is similar (same text style)
            if y_distance <= y_tolerance and font_diff <= font_size_tolerance:
                # Add to current row
                current_row['items'].append(item)
                current_row['y_center'] = (current_row['y_center'] + item_y_center) / 2
                current_row['font_size'] = (current_row['font_size'] + item.font_size) / 2
                current_row['y_min'] = min(current_row['y_min'], item.y0)
                current_row['y_max'] = max(current_row['y_max'], item.y1)
            else:
                # Save current row and start new one
                rows.append(current_row)
                current_row = {
                    'items': [item],
                    'y_center': item_y_center,
                    'font_size': item.font_size,
                    'y_min': item.y0,
                    'y_max': item.y1
                }
    
    # Add last row
    if current_row:
        rows.append(current_row)
    
    return rows


def merge_multiline_cells(cells: List[Cell], y_gap_threshold: float = 5.0) -> List[Cell]:
    """
    Merge cells that are in same column with small Y gap (multi-line text).
    
    Args:
        cells: List of cells
        y_gap_threshold: Maximum Y gap for merging
        
    Returns:
        List of merged cells
    """
    # Group cells by column
    cells_by_col = defaultdict(list)
    for cell in cells:
        cells_by_col[cell.col_index].append(cell)
    
    merged_cells = []
    
    for col_idx, col_cells in cells_by_col.items():
        # Sort by row index
        col_cells.sort(key=lambda x: x.row_index)
        
        current_cell = None
        for cell in col_cells:
            if current_cell is None:
                current_cell = cell
            else:
                # Check if cells should be merged (same column, small Y gap)
                y_gap = abs(cell.y0 - current_cell.y1)
                
                if y_gap <= y_gap_threshold:
                    # Merge cells
                    current_cell.text += " " + cell.text
                    current_cell.y1 = max(current_cell.y1, cell.y1)
                    current_cell.y0 = min(current_cell.y0, cell.y0)
                else:
                    # Save current cell and start new one
                    merged_cells.append(current_cell)
                    current_cell = cell
        
        if current_cell:
            merged_cells.append(current_cell)
    
    return merged_cells


def detect_header_footer(text_items: List[TextItem], all_pages: List[int]) -> Tuple[List[str], List[str]]:
    """
    Detect repeating text across pages (headers/footers).
    
    Args:
        text_items: All text items
        all_pages: List of page numbers
        
    Returns:
        Tuple of (header_texts, footer_texts) to suppress
    """
    if len(all_pages) < 2:
        return [], []
    
    # Get items from top 10% and bottom 10% of each page
    headers = []
    footers = []
    
    for page_num in all_pages:
        page_items = [item for item in text_items if item.page_num == page_num]
        if not page_items:
            continue
        
        # Get page height range
        y_max = max(item.y1 for item in page_items)
        y_min = min(item.y0 for item in page_items)
        page_height = y_max - y_min
        
        header_threshold = y_max - (page_height * 0.1)  # Top 10%
        footer_threshold = y_min + (page_height * 0.1)  # Bottom 10%
        
        # Collect header and footer texts
        for item in page_items:
            if item.y1 >= header_threshold:
                headers.append(item.text.strip())
            elif item.y0 <= footer_threshold:
                footers.append(item.text.strip())
    
    # Find texts that appear in >50% of pages (likely headers/footers)
    from collections import Counter
    header_counts = Counter(headers)
    footer_counts = Counter(footers)
    
    min_occurrences = len(all_pages) * 0.5
    header_texts = [text for text, count in header_counts.items() if count >= min_occurrences]
    footer_texts = [text for text, count in footer_counts.items() if count >= min_occurrences]
    
    return header_texts, footer_texts


def detect_visual_pdf(pdf_path: str) -> Tuple[bool, str]:
    """
    Detect if PDF is visual/design-based (high image density, boxes, etc.).
    
    Args:
        pdf_path: Path to PDF file
        
    Returns:
        Tuple of (is_visual, reason)
    """
    try:
        from pdfminer.layout import LTFigure, LTImage
        import statistics
        
        text_items = extract_text_with_positions(pdf_path)
        
        if not text_items:
            return True, "No extractable text found"
        
        # Get page layouts for image/box detection
        laparams = LAParams()
        pages = list(extract_pages(pdf_path, laparams=laparams))
        
        total_text_area = 0
        total_image_area = 0
        box_count = 0
        
        for page_layout in pages:
            page_width = page_layout.width
            page_height = page_layout.height
            page_area = page_width * page_height
            
            text_area = 0
            image_area = 0
            
            for element in page_layout:
                # Check for images
                if isinstance(element, LTFigure):
                    box_count += 1
                    x0, y0, x1, y1 = element.bbox
                    area = (x1 - x0) * (y1 - y0)
                    image_area += area
                    
                    # Check for images inside figure
                    for child in element:
                        if isinstance(child, LTImage):
                            x0, y0, x1, y1 = child.bbox
                            area = (x1 - x0) * (y1 - y0)
                            image_area += area
                
                # Calculate text area
                if isinstance(element, LTTextBox):
                    x0, y0, x1, y1 = element.bbox
                    text_area += (x1 - x0) * (y1 - y0)
            
            total_text_area += text_area
            total_image_area += image_area
        
        # Calculate ratios
        total_area = sum(p.width * p.height for p in pages)
        image_ratio = total_image_area / total_area if total_area > 0 else 0
        text_ratio = total_text_area / total_area if total_area > 0 else 0
        
        # Detection criteria
        if image_ratio > 0.3:  # More than 30% images
            return True, "High image density detected (>30%)"
        
        if text_ratio < 0.1:  # Less than 10% text
            return True, "Very low text content (<10%)"
        
        if box_count > len(pages) * 5:  # More than 5 boxes per page average
            return True, f"High box/rectangle count ({box_count} boxes)"
        
        # Check for rectangular layouts (many items with similar x positions indicating columns/boxes)
        if len(text_items) > 100:
            x_positions = [item.x0 for item in text_items]
            # Check if many items align to same X positions (rectangular layout)
            x_buckets = defaultdict(int)
            for x in x_positions:
                x_bucket = round(x / 10) * 10  # Round to nearest 10
                x_buckets[x_bucket] += 1
            
            # If many items align to same X positions, it might be a structured/form layout
            max_bucket_count = max(x_buckets.values()) if x_buckets else 0
            if max_bucket_count > len(text_items) * 0.3:
                return True, "Rectangular/structured layout detected (form-like)"
        
        return False, ""
    
    except Exception as e:
        logger.warning(f"Error detecting visual PDF: {e}")
        # On error, allow conversion (don't block)
        return False, ""


def convert_pdf_to_docx_free_pipeline(pdf_path: str, output_path: str) -> str:
    """
    Convert PDF to DOCX using FREE pipeline with data-oriented layout reconstruction.
    
    Features:
    - X-axis column detection
    - Improved row detection
    - Multi-line cell merge
    - Page-wise isolation
    - Header/footer suppression
    
    Args:
        pdf_path: Path to input PDF
        output_path: Path to output DOCX
        
    Returns:
        Path to created DOCX file
        
    Raises:
        ValueError: If PDF is detected as visual/design-based
    """
    logger.info(f"Converting PDF to DOCX with FREE pipeline: {pdf_path}")
    
    # Detect visual PDF (block if detected)
    is_visual, reason = detect_visual_pdf(pdf_path)
    if is_visual:
        error_msg = f"This document contains complex visual layout. {reason} Use Premium for accurate Excel conversion."
        logger.warning(f"Visual PDF detected, blocking FREE conversion: {reason}")
        raise ValueError(error_msg)
    
    # Extract text with positions
    text_items = extract_text_with_positions(pdf_path)
    
    if not text_items:
        raise ValueError("No text content found in PDF")
    
    # Get unique pages
    all_pages = sorted(set(item.page_num for item in text_items))
    
    # Detect headers/footers
    header_texts, footer_texts = detect_header_footer(text_items, all_pages)
    logger.info(f"Detected {len(header_texts)} header texts and {len(footer_texts)} footer texts to suppress")
    
    # Create Word document
    doc = Document()
    
    # Process each page separately (page-wise isolation)
    for page_num in all_pages:
        if page_num > 1:
            doc.add_page_break()
        
        # Get page items, excluding headers/footers
        page_items = [
            item for item in text_items
            if item.page_num == page_num
            and item.text.strip() not in header_texts
            and item.text.strip() not in footer_texts
        ]
        
        if not page_items:
            continue
        
        # Detect columns for this page
        columns = detect_columns(page_items, page_num)
        logger.debug(f"Page {page_num}: Detected {len(columns)} columns")
        
        # Detect rows for this page
        rows = detect_rows(page_items, page_num)
        logger.debug(f"Page {page_num}: Detected {len(rows)} rows")
        
        if not columns or not rows:
            # Fallback: Just add paragraphs
            for item in sorted(page_items, key=lambda x: x.y0, reverse=True):
                doc.add_paragraph(item.text)
            continue
        
        # Create cells: assign each text item to row and column
        cells = []
        for row_idx, row in enumerate(rows):
            for item in row['items']:
                col_idx = assign_column(item, columns)
                if col_idx is not None:
                    cells.append(Cell(
                        text=item.text,
                        col_index=col_idx,
                        row_index=row_idx,
                        x0=item.x0,
                        x1=item.x1,
                        y0=item.y0,
                        y1=item.y1,
                        font_size=item.font_size
                    ))
        
        # Merge multi-line cells
        merged_cells = merge_multiline_cells(cells)
        
        # Determine table dimensions
        if merged_cells:
            max_col = max(cell.col_index for cell in merged_cells)
            max_row = max(cell.row_index for cell in merged_cells)
            
            # Create table in Word
            table = doc.add_table(rows=max_row + 1, cols=max_col + 1)
            
            # Fill table cells
            for cell in merged_cells:
                if cell.row_index <= max_row and cell.col_index <= max_col:
                    word_cell = table.rows[cell.row_index].cells[cell.col_index]
                    word_cell.text = cell.text
            
            # Add spacing after table
            doc.add_paragraph("")
        else:
            # No table structure detected, add as paragraphs
            for item in sorted(page_items, key=lambda x: x.y0, reverse=True):
                doc.add_paragraph(item.text)
    
    # Save document
    import os
    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
    doc.save(output_path)
    
    logger.info(f"FREE pipeline conversion successful: {output_path}")
    return output_path

