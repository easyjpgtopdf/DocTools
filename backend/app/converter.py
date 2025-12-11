
"""
PDF to Word conversion module.
Handles PDF text detection and conversion using LibreOffice or Document AI.
"""
import subprocess
import os
import time
import logging
from typing import Tuple, Optional, List
from pathlib import Path
from enum import Enum
from dataclasses import dataclass
from pdfminer.high_level import extract_text as pdf_extract_text
from pdfminer.layout import LAParams
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.docai_client import process_pdf_to_layout, ParsedDocument, get_docai_confidence

logger = logging.getLogger(__name__)

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    logger.warning("PyMuPDF (fitz) not available. PDF to JPG conversion will not work.")

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False
    logger.warning("Pillow (PIL) not available. Image processing will be limited.")


class ConversionMethod(str, Enum):
    """Enum for conversion methods."""
    LIBREOFFICE = "libreoffice"
    DOCAI = "docai"


@dataclass
class SmartConversionResult:
    """Result of smart PDF to Word conversion with primary and alternative."""
    main_docx_path: str
    main_method: ConversionMethod
    alt_docx_path: Optional[str]
    alt_method: Optional[ConversionMethod]
    pages: int
    used_docai: bool
    conversion_time_ms: int
    docai_confidence: Optional[float] = None  # Average confidence from DocAI if used
    parsed_document: Optional[ParsedDocument] = None  # Store parsed doc for confidence extraction


def pdf_has_text(pdf_path: str, max_pages_to_check: int = 3) -> bool:
    """
    Check if PDF contains extractable text.
    
    Args:
        pdf_path: Path to PDF file
        max_pages_to_check: Maximum number of pages to check
        
    Returns:
        True if PDF has extractable text, False otherwise
    """
    try:
        logger.info(f"Checking if PDF has text: {pdf_path}")
        
        # Extract text from first few pages
        laparams = LAParams()
        text = pdf_extract_text(
            pdf_path,
            maxpages=max_pages_to_check,
            laparams=laparams
        )
        
        # Clean and check if meaningful text exists
        cleaned_text = text.strip().replace('\n', '').replace(' ', '')
        
        has_text = len(cleaned_text) > 50  # At least 50 characters
        
        logger.info(f"PDF has text: {has_text} (found {len(cleaned_text)} characters)")
        return has_text
        
    except Exception as e:
        logger.error(f"Error checking PDF text: {e}")
        # If we can't determine, assume it doesn't have text (safer to use DocAI)
        return False


def convert_pdf_to_docx_with_libreoffice(pdf_path: str, output_dir: str) -> str:
    """
    Convert PDF to DOCX using LibreOffice headless (soffice).
    
    This function uses LibreOffice's command-line interface to convert
    PDF files with extractable text to DOCX format. LibreOffice is ideal
    for PDFs that already contain text layers.
    
    Example subprocess command executed:
        soffice --headless --convert-to docx --outdir /tmp input.pdf
    
    Implementation details:
        - Uses subprocess.run() to execute soffice command
        - Runs in headless mode (no GUI)
        - Converts PDF to DOCX format
        - Saves output to specified directory
        - Handles timeouts (5 minutes) and errors gracefully
    
    Args:
        pdf_path: Path to input PDF file
        output_dir: Directory to save output DOCX
        
    Returns:
        Path to created DOCX file
        
    Raises:
        Exception: If conversion fails or times out
    """
    try:
        logger.info(f"Converting PDF to DOCX with LibreOffice: {pdf_path}")
        
        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)
        
        # Check for LibreOffice binary in common locations
        soffice_paths = [
            "/usr/bin/libreoffice",
            "/usr/bin/soffice",
            "soffice",  # Try PATH as fallback
        ]
        
        soffice_binary = None
        for path in soffice_paths:
            if os.path.exists(path) or path == "soffice":
                # Check if it's executable (for absolute paths) or in PATH
                if path != "soffice":
                    if os.access(path, os.X_OK):
                        soffice_binary = path
                        break
                else:
                    # Check if soffice is in PATH
                    try:
                        result = subprocess.run(
                            ["which", "soffice"],
                            capture_output=True,
                            timeout=2
                        )
                        if result.returncode == 0:
                            soffice_binary = "soffice"
                            break
                    except:
                        pass
        
        if not soffice_binary:
            logger.error("LibreOffice (soffice) not found in /usr/bin/libreoffice, /usr/bin/soffice, or PATH")
            raise Exception("LibreOffice not available in Cloud Run. Please check installation.")
        
        logger.info(f"Using LibreOffice binary: {soffice_binary}")
        
        # Build LibreOffice command
        # --headless: Run without GUI
        # --convert-to docx: Convert to DOCX format
        # --outdir: Output directory
        cmd = [
            soffice_binary,
            "--headless",           # No GUI mode
            "--convert-to", "docx", # Target format
            "--outdir", output_dir, # Output directory
            pdf_path                # Input file
        ]
        
        logger.debug(f"Executing command: {' '.join(cmd)}")
        print(f"DEBUG: LibreOffice command: {' '.join(cmd)}")
        
        # Run LibreOffice with timeout protection
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=300  # 5 minute timeout for large files
        )
        
        logger.debug(f"LibreOffice return code: {result.returncode}")
        logger.debug(f"LibreOffice stdout: {result.stdout}")
        logger.debug(f"LibreOffice stderr: {result.stderr}")
        print(f"DEBUG: LibreOffice return code: {result.returncode}")
        print(f"DEBUG: LibreOffice stdout: {result.stdout}")
        print(f"DEBUG: LibreOffice stderr: {result.stderr}")
        
        # Check if command succeeded
        if result.returncode != 0:
            error_msg = result.stderr or result.stdout or "Unknown error"
            logger.error(f"LibreOffice conversion failed: {error_msg}")
            raise Exception(f"LibreOffice conversion failed: {error_msg}")
        
        # LibreOffice creates output with same base name as input
        # e.g., input.pdf -> input.docx
        # But LibreOffice might keep original filename extension or use different naming
        pdf_name = Path(pdf_path).stem
        pdf_basename = Path(pdf_path).name
        
        # Wait a moment for LibreOffice to finish writing the file
        import time
        time.sleep(0.5)
        
        # List all files in output directory
        output_files = []
        if os.path.exists(output_dir):
            try:
                output_files = os.listdir(output_dir)
            except Exception as e:
                logger.warning(f"Error listing output directory: {e}")
        
        logger.info(f"Files in output directory after conversion: {output_files}")
        print(f"DEBUG: Files in output directory: {output_files}")
        
        # Check for DOCX file - LibreOffice may use different naming conventions
        # Try multiple possible output names
        possible_names = [
            f"{pdf_name}.docx",  # input -> input.docx
            pdf_basename.replace('.pdf', '.docx'),  # input.pdf -> input.docx
            f"{pdf_basename}.docx",  # input.pdf -> input.pdf.docx
            "output.docx",  # Generic output name
        ]
        
        actual_output_path = None
        
        # First, try the expected paths
        for name in possible_names:
            test_path = os.path.join(output_dir, name)
            if os.path.exists(test_path) and os.path.getsize(test_path) > 0:
                actual_output_path = test_path
                logger.info(f"Found DOCX at expected path: {actual_output_path}")
                break
        
        # If not found, search for any .docx file in the directory
        if not actual_output_path:
            for file in output_files:
                if file.endswith('.docx') and not file.startswith('~'):  # Exclude temp files
                    test_path = os.path.join(output_dir, file)
                    if os.path.exists(test_path) and os.path.getsize(test_path) > 0:
                        actual_output_path = test_path
                        logger.info(f"Found DOCX file with different name: {actual_output_path}")
                        break
        
        # If still not found, check if LibreOffice created it in a subdirectory or parent
        if not actual_output_path:
            # Check parent directory
            parent_dir = os.path.dirname(output_dir)
            if os.path.exists(parent_dir):
                try:
                    parent_files = os.listdir(parent_dir)
                    for file in parent_files:
                        if file.endswith('.docx'):
                            test_path = os.path.join(parent_dir, file)
                            if os.path.exists(test_path):
                                actual_output_path = test_path
                                logger.info(f"Found DOCX in parent directory: {actual_output_path}")
                                break
                except Exception:
                    pass
        
        # Final verification
        if not actual_output_path or not os.path.exists(actual_output_path):
            error_msg = f"Output DOCX not found. Expected at: {os.path.join(output_dir, f'{pdf_name}.docx')}. Files in directory: {output_files}. LibreOffice may have failed silently."
            logger.error(error_msg)
            raise Exception(error_msg)
        
        # Verify file is not empty
        if os.path.getsize(actual_output_path) == 0:
            raise Exception(f"Output DOCX file is empty: {actual_output_path}")
        
        logger.info(f"Conversion successful: {actual_output_path} (size: {os.path.getsize(actual_output_path)} bytes)")
        print(f"DEBUG: LibreOffice output: {actual_output_path}")
        return actual_output_path
        
    except subprocess.TimeoutExpired:
        logger.error("LibreOffice conversion timed out after 5 minutes", exc_info=True)
        raise Exception("Conversion timed out. File may be too large or complex.")
    except FileNotFoundError as e:
        logger.error(f"LibreOffice (soffice) not found in PATH: {e}", exc_info=True)
        raise Exception("LibreOffice not available in Cloud Run. Please check installation.")
    except Exception as e:
        logger.error(f"LibreOffice conversion error: {e}", exc_info=True)
        raise


def convert_pdf_to_docx_with_docai(
    pdf_path: str,
    output_path: str,
    parsed_doc: ParsedDocument
) -> str:
    """
    Convert PDF to DOCX using Document AI parsed content.
    
    This function takes structured content extracted by Document AI (from scanned PDFs)
    and creates a Word document using python-docx. It preserves:
    - Page structure (page breaks)
    - Text paragraphs
    - Headings (based on font size detection)
    - Tables (if detected by Document AI)
    
    Example Document AI loop over paragraphs:
        for page_num, page_blocks in enumerate(parsed_doc.pages, start=1):
            for block in page_blocks:
                if block.is_heading:
                    doc.add_heading(block.text, level=2)
                else:
                    doc.add_paragraph(block.text)
    
    Implementation details:
        - Creates a new Document() using python-docx
        - Iterates over pages and text blocks from Document AI
        - Preserves headings based on font size heuristics
        - Extracts and formats tables if present
        - Saves the final DOCX document
    
    Args:
        pdf_path: Path to input PDF file (for reference only)
        output_path: Path to save output DOCX file
        parsed_doc: ParsedDocument object from Document AI containing structured content
        
    Returns:
        Path to created DOCX file
    """
    try:
        logger.info(f"Converting PDF to DOCX with Document AI: {pdf_path}")
        
        # Create new Word document
        doc = Document()
        
        # Process each page
        for page_num, page_blocks in enumerate(parsed_doc.pages, start=1):
            # Add page break before each page (except first)
            if page_num > 1:
                doc.add_page_break()
            
            # Process text blocks
            for block in page_blocks:
                if block.is_heading:
                    # Determine heading level based on font size
                    heading_level = 3
                    if block.font_size and block.font_size >= 18:
                        heading_level = 1
                    elif block.font_size and block.font_size >= 14:
                        heading_level = 2
                    
                    heading = doc.add_heading(block.text, level=heading_level)
                    
                    # Apply font to heading if available
                    if block.font_name:
                        try:
                            for run in heading.runs:
                                # Map Hindi fonts to Unicode-compatible fonts
                                hindi_font_map = {
                                    'krutidev': 'Noto Sans Devanagari',
                                    'krutidev010': 'Noto Sans Devanagari',
                                    'chanakya': 'Noto Sans Devanagari',
                                    'mangal': 'Noto Sans Devanagari',
                                }
                                font_lower = block.font_name.lower()
                                if any(hindi_font in font_lower for hindi_font in hindi_font_map.keys()):
                                    run.font.name = hindi_font_map.get(font_lower.split('-')[0], 'Noto Sans Devanagari')
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
                                else:
                                    run.font.name = block.font_name
                        except Exception as e:
                            logger.warning(f"Could not set heading font: {e}")
                    
                    # Auto-detect Hindi in headings
                    if not block.font_name and any(ord(char) >= 0x0900 and ord(char) <= 0x097F for char in block.text):
                        try:
                            for run in heading.runs:
                                run.font.name = 'Noto Sans Devanagari'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Devanagari')
                        except Exception:
                            pass
                else:
                    # Regular paragraph with Unicode and font preservation
                    p = doc.add_paragraph()
                    run = p.add_run(block.text)
                    
                    # Set font size if available
                    if block.font_size:
                        run.font.size = Pt(block.font_size)
                    
                    # Preserve font name if available (for Hindi fonts like Krutidev, Chanakya)
                    if block.font_name:
                        try:
                            # Map common font names to supported fonts
                            font_name = block.font_name
                            
                            # Hindi font mappings
                            hindi_font_map = {
                                'krutidev': 'Noto Sans Devanagari',
                                'krutidev010': 'Noto Sans Devanagari',
                                'chanakya': 'Noto Sans Devanagari',
                                'mangal': 'Noto Sans Devanagari',
                                'aaparajita': 'Noto Sans Devanagari',
                            }
                            
                            # Convert to lowercase for matching
                            font_lower = font_name.lower()
                            
                            # Check if it's a Hindi font
                            if any(hindi_font in font_lower for hindi_font in hindi_font_map.keys()):
                                # Use Unicode-compatible Hindi font
                                run.font.name = hindi_font_map.get(font_lower.split('-')[0], 'Noto Sans Devanagari')
                                # Set font for complex scripts (Hindi, Devanagari)
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
                            else:
                                # Try to use original font name
                                run.font.name = font_name
                        except Exception as e:
                            logger.warning(f"Could not set font '{block.font_name}': {e}. Using default.")
                            # Default to Noto Sans Devanagari for Hindi/Devanagari text
                            if any(ord(char) >= 0x0900 and ord(char) <= 0x097F for char in block.text):
                                run.font.name = 'Noto Sans Devanagari'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Devanagari')
                    
                    # Auto-detect Hindi/Devanagari Unicode range and set appropriate font
                    if not block.font_name:
                        # Check if text contains Devanagari characters (Unicode range 0x0900-0x097F)
                        if any(ord(char) >= 0x0900 and ord(char) <= 0x097F for char in block.text):
                            try:
                                run.font.name = 'Noto Sans Devanagari'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Devanagari')
                                logger.debug(f"Auto-detected Hindi/Devanagari text, using Noto Sans Devanagari font")
                            except Exception as e:
                                logger.warning(f"Could not set Hindi font: {e}")
                    
                    # Ensure Unicode encoding is preserved
                    # python-docx handles Unicode automatically, but we ensure proper encoding
                    try:
                        # Verify text is properly encoded
                        block.text.encode('utf-8')
                    except UnicodeEncodeError:
                        logger.warning(f"Unicode encoding issue with text: {block.text[:50]}...")
        
        # Add tables if present
        for table_data in parsed_doc.tables:
            # Add page break if table is on different page
            if table_data.page_number > 1:
                doc.add_page_break()
            
            # Create table in Word document
            table = doc.add_table(rows=len(table_data.rows), cols=len(table_data.rows[0]) if table_data.rows else 0)
            table.style = 'Light Grid Accent 1'
            
            # Populate table cells
            for row_idx, row_data in enumerate(table_data.rows):
                for col_idx, cell_data in enumerate(row_data):
                    if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                        table.rows[row_idx].cells[col_idx].text = str(cell_data)
        
        # Save document
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        doc.save(output_path)
        
        logger.info(f"Document AI conversion successful: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"Document AI conversion error: {e}")
        raise


def get_page_count(pdf_path: str) -> int:
    """
    Get page count from PDF file.
    
    Args:
        pdf_path: Path to PDF file
        
    Returns:
        Number of pages in the PDF
    """
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(pdf_path)
        return len(reader.pages)
    except ImportError:
        # Fallback to pdfminer if PyPDF2 is not available
        try:
            from pdfminer.high_level import extract_pages
            pages = list(extract_pages(pdf_path))
            return len(pages)
        except Exception as e:
            logger.warning(f"Could not determine page count: {e}")
            return 0
    except Exception as e:
        logger.warning(f"Error getting page count: {e}")
        return 0


def convert_pdf_to_docx(
    pdf_path: str,
    output_dir: str,
    use_docai: bool = False,
    parsed_doc: Optional[ParsedDocument] = None
) -> Tuple[str, bool]:
    """
    Main conversion function. Chooses appropriate method.
    
    Args:
        pdf_path: Path to input PDF
        output_dir: Directory for output DOCX
        use_docai: Force use of Document AI
        parsed_doc: Pre-parsed Document AI result (optional)
        
    Returns:
        Tuple of (output_path, used_docai)
    """
    os.makedirs(output_dir, exist_ok=True)
    pdf_name = Path(pdf_path).stem
    
    if use_docai or parsed_doc:
        # Use Document AI path
        if not parsed_doc:
            raise ValueError("ParsedDocument required when use_docai=True")
        
        output_path = os.path.join(output_dir, f"{pdf_name}.docx")
        convert_pdf_to_docx_with_docai(pdf_path, output_path, parsed_doc)
        return output_path, True
    else:
        # Use LibreOffice path
        output_path = convert_pdf_to_docx_with_libreoffice(pdf_path, output_dir)
        return output_path, False


def smart_convert_pdf_to_word(
    pdf_path: str,
    output_dir: str,
    settings,
    generate_alternative: bool = True
) -> SmartConversionResult:
    """
    Intelligently convert PDF to Word, choosing best method and optionally generating alternative.
    
    This function automatically detects if the PDF has extractable text:
    - If text exists: Uses LibreOffice (primary), Document AI (alternative if requested)
    - If no text (scanned): Uses Document AI (primary), LibreOffice (alternative if requested)
    
    Args:
        pdf_path: Path to input PDF file
        output_dir: Directory for output DOCX files
        settings: Settings object with project_id, docai_location, docai_processor_id
        generate_alternative: Whether to generate alternative conversion
        
    Returns:
        SmartConversionResult with primary and optional alternative conversions
    """
    start_time = time.time()
    
    # Check if PDF has text
    has_text = pdf_has_text(pdf_path)
    pages = get_page_count(pdf_path)
    
    # Determine primary method
    if has_text:
        primary_method = ConversionMethod.LIBREOFFICE
        logger.info("Using LibreOffice as primary method (text-based PDF)")
    else:
        primary_method = ConversionMethod.DOCAI
        logger.info("Using Document AI as primary method (scanned PDF)")
    
    # Initialize variables
    docai_confidence = None
    parsed_doc = None
    primary_docx = None
    alt_docx = None
    alt_method = None
    
    # Generate primary conversion
    primary_start = time.time()
    if primary_method == ConversionMethod.LIBREOFFICE:
        primary_docx = convert_pdf_to_docx_with_libreoffice(pdf_path, output_dir)
        used_docai = False
    else:
        # DocAI conversion
        try:
            # Validate Document AI settings
            if not settings.project_id:
                raise Exception("Document AI: PROJECT_ID is required")
            if not settings.docai_processor_id:
                raise Exception("Document AI: DOCAI_PROCESSOR_ID is required")
            if not settings.docai_location:
                settings.docai_location = "us"  # Default location
            
            logger.info(f"Using Document AI: project_id={settings.project_id}, processor_id={settings.docai_processor_id}, location={settings.docai_location}")
            
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
            
            parsed_doc = process_pdf_to_layout(
                settings.project_id,
                settings.docai_location,
                settings.docai_processor_id,
                pdf_bytes
            )
            
            # Extract DocAI confidence
            docai_confidence = get_docai_confidence(parsed_doc)
            if docai_confidence is not None:
                logger.info(f"DocAI confidence extracted: {docai_confidence:.3f}")
            
            primary_docx = os.path.join(output_dir, f"primary-{Path(pdf_path).stem}.docx")
            convert_pdf_to_docx_with_docai(pdf_path, primary_docx, parsed_doc)
            used_docai = True
        except Exception as docai_error:
            logger.error(f"Document AI conversion failed: {docai_error}", exc_info=True)
            # Fallback to basic text extraction if DocAI fails
            logger.warning("Falling back to basic text extraction due to Document AI failure")
            # For now, re-raise the error - in future we can implement basic fallback
            raise Exception(f"Document AI failed: {str(docai_error)}. Please check DOCAI_PROCESSOR_ID, project_id, and credentials.")
    
    primary_time = time.time() - primary_start
    
    # Generate alternative conversion if requested
    if generate_alternative:
        alt_start = time.time()
        if primary_method == ConversionMethod.LIBREOFFICE:
            # Alternative: Use DocAI
            alt_method = ConversionMethod.DOCAI
            logger.info("Generating alternative using Document AI")
            try:
                with open(pdf_path, "rb") as f:
                    pdf_bytes = f.read()
                
                parsed_doc_alt = process_pdf_to_layout(
                    settings.project_id,
                    settings.docai_location,
                    settings.docai_processor_id,
                    pdf_bytes
                )
                
                # Use parsed doc for confidence if primary didn't use DocAI
                if docai_confidence is None:
                    docai_confidence = get_docai_confidence(parsed_doc_alt)
                    if docai_confidence is not None:
                        logger.info(f"DocAI confidence extracted from alternative: {docai_confidence:.3f}")
                
                alt_docx = os.path.join(output_dir, f"alt-{Path(pdf_path).stem}.docx")
                convert_pdf_to_docx_with_docai(pdf_path, alt_docx, parsed_doc_alt)
            except Exception as alt_error:
                logger.warning(f"Alternative DocAI conversion failed: {alt_error}. Skipping alternative.")
                alt_docx = None
                alt_method = None
        else:
            # Alternative: Use LibreOffice
            alt_method = ConversionMethod.LIBREOFFICE
            logger.info("Generating alternative using LibreOffice")
            alt_docx = convert_pdf_to_docx_with_libreoffice(pdf_path, output_dir)
        
        alt_time = time.time() - alt_start
        logger.info(f"Alternative conversion completed in {alt_time:.2f}s")
    
    total_time_ms = int((time.time() - start_time) * 1000)
    
    result = SmartConversionResult(
        main_docx_path=primary_docx,
        main_method=primary_method,
        alt_docx_path=alt_docx,
        alt_method=alt_method,
        pages=pages,
        used_docai=used_docai,
        conversion_time_ms=total_time_ms,
        docai_confidence=docai_confidence,
        parsed_document=parsed_doc if used_docai else None
    )
    
    logger.info(f"Smart conversion completed in {total_time_ms}ms")
    return result


# PDF to JPG conversion (using PyMuPDF)
if HAS_PYMUPDF:
    def convert_pdf_to_jpg(pdf_path: str, output_dir: str, dpi: int = 150) -> List[str]:
        """
        Convert each page of a PDF to a JPG image.
        
        Args:
            pdf_path: Path to input PDF file
            output_dir: Directory to save JPG images
            dpi: Resolution for output images (default: 150)
            
        Returns:
            List of paths to created JPG files
        """
        logger.info(f"Converting PDF to JPG: {pdf_path} with {dpi} DPI")
        output_image_paths = []
        
        try:
            doc = fitz.open(pdf_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap(matrix=fitz.Matrix(dpi / 72, dpi / 72))
                
                output_filename = f"page_{page_num + 1:03d}.jpg"
                output_path = os.path.join(output_dir, output_filename)
                pix.save(output_path)
                output_image_paths.append(output_path)
                logger.debug(f"Saved JPG for page {page_num + 1} to {output_path}")
            
            doc.close()
            logger.info(f"Converted {len(output_image_paths)} pages to JPG.")
            return output_image_paths
        except Exception as e:
            logger.error(f"Error converting PDF to JPG: {e}")
            raise
else:
    def convert_pdf_to_jpg(pdf_path: str, output_dir: str, dpi: int = 150) -> List[str]:
        raise ImportError("PyMuPDF (fitz) is required for PDF to JPG conversion.")


# PDF thumbnail generation (using PyMuPDF)
if HAS_PYMUPDF:
    def generate_pdf_thumbnail(pdf_path: str, output_path: str, width_px: int = 300) -> str:
        """
        Generate a thumbnail (first page as JPG) of a PDF.
        
        Args:
            pdf_path: Path to input PDF file
            output_path: Path to save thumbnail JPG
            width_px: Desired width of thumbnail in pixels (default: 300)
            
        Returns:
            Path to created thumbnail file
        """
        logger.info(f"Generating thumbnail for {pdf_path} with width {width_px}px")
        
        try:
            doc = fitz.open(pdf_path)
            page = doc.load_page(0)  # First page
            
            # Calculate DPI based on desired width
            zoom = width_px / page.rect.width
            mat = fitz.Matrix(zoom, zoom)
            
            pix = page.get_pixmap(matrix=mat)
            pix.save(output_path)
            doc.close()
            
            logger.info(f"Thumbnail saved to {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Error generating PDF thumbnail: {e}")
            raise
else:
    def generate_pdf_thumbnail(pdf_path: str, output_path: str, width_px: int = 300) -> str:
        raise ImportError("PyMuPDF (fitz) is required for PDF thumbnail generation.")
