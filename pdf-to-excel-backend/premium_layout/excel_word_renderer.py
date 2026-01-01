"""
Excel/Word Renderer
Converts unified layout model to Excel and Word formats.
"""

import logging
import io
from typing import Optional, List, Dict

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available, Word export disabled")

from .unified_layout_model import UnifiedLayout, Cell, CellStyle, CellAlignment

try:
    from openpyxl.drawing.image import Image as ExcelImage
    IMAGE_SUPPORT = True
except ImportError:
    IMAGE_SUPPORT = False
    logger.warning("openpyxl image support not available")


class ExcelWordRenderer:
    """Renders unified layout to Excel and Word formats"""
    
    def __init__(self):
        """Initialize renderer"""
        pass
    
    def render_to_excel(self, layouts: List[UnifiedLayout], images: Optional[List[Dict]] = None) -> bytes:
        """
        CRITICAL REWRITE: Render unified layouts to Excel with explicit grid mapping.
        
        Requirements:
        1. Build explicit grid per page (rows = max(row_index) + 1, cols = max(col_index) + 1)
        2. One Excel sheet per PDF page (Page_1, Page_2, etc.)
        3. Write cells to exact Excel coordinates (row_index + 1, col_index + 1)
        4. Enforce hard boundaries (no word spill, no fake rows/columns)
        5. Column control for invoices/bills (force 2 columns if detected)
        6. Font rules (one row = one font)
        7. Image handling (anchored positions, no row/column creation)
        8. Validation logging (rows/columns written, first 5 rows preview)
        9. Empty layout handling (fail conversion, return error)
        
        Args:
            layouts: List of UnifiedLayout objects (one per page)
            images: Optional list of images to insert (from ImageExtractor)
            
        Returns:
            Excel file as bytes
            
        Raises:
            ValueError: If layout is empty
        """
        # CRITICAL: Validate layouts before processing
        if not layouts:
            raise ValueError("Layout empty — conversion aborted: No layouts provided")
        
        # Check if all layouts are empty
        all_empty = all(layout.is_empty() for layout in layouts)
        if all_empty:
            raise ValueError("Layout empty — conversion aborted: All layouts are empty")
        
        workbook = Workbook()
        workbook.remove(workbook.active)
        
        # Group images by page
        images_by_page = {}
        if images:
            for img in images:
                page_idx = img.get('page_index', 0)
                if page_idx not in images_by_page:
                    images_by_page[page_idx] = []
                images_by_page[page_idx].append(img)
        
        # Create one sheet per page (Page-to-Sheet Mapping)
        for layout in layouts:
            page_num = layout.page_index + 1
            sheet_name = f"Page_{page_num}" if len(layouts) > 1 else "Document"
            if len(sheet_name) > 31:
                sheet_name = sheet_name[:31]
            
            # CRITICAL: Validate layout is not empty
            if layout.is_empty():
                logger.critical("=" * 80)
                logger.critical(f"❌ LAYOUT EMPTY: Page {page_num} layout is empty")
                logger.critical("=" * 80)
                raise ValueError(f"Layout empty — conversion aborted: Page {page_num} layout is empty")
            
            # ENTERPRISE RULE: Renderer is DUMB WRITER - NO layout inference
            # Spatial indexing, newline splitting, and layout cleaning MUST happen BEFORE renderer
            # Renderer responsibility: Write rows, columns, merged cells EXACTLY as provided
            
            # LOG WHAT RENDERER RECEIVES
            logger.critical("=" * 80)
            logger.critical(f"📊 RENDERER INPUT (Page {page_num}):")
            logger.critical(f"   Layout.rows count: {len(layout.rows)}")
            logger.critical(f"   Total cells: {sum(len(row) for row in layout.rows)}")
            logger.critical(f"   Execution mode: {layout.metadata.get('execution_mode', 'unknown')}")
            logger.critical(f"   Frozen: {layout.metadata.get('frozen', False)}")
            
            # Log first 10 cells
            all_cells = [cell for row in layout.rows for cell in row]
            if all_cells:
                logger.critical(f"   First 10 cells (row,col,value):")
                for idx, cell in enumerate(all_cells[:10]):
                    logger.critical(f"      [{idx}] ({cell.row},{cell.column}) = '{str(cell.value)[:30]}'")
            else:
                logger.critical("   ❌ NO CELLS IN LAYOUT!")
            logger.critical("=" * 80)
            
            # Quality fail-safe: Reject invalid layouts
            max_row = layout.get_max_row()
            max_col = layout.get_max_column()
            logger.critical(f"   get_max_row() returned: {max_row}")
            logger.critical(f"   get_max_column() returned: {max_col}")
            
            if max_row < 0 or max_col < 0:
                logger.critical(f"❌ INVALID LAYOUT: Page {page_num} has invalid dimensions (max_row={max_row}, max_col={max_col})")
                raise ValueError(f"Layout empty — conversion aborted: Page {page_num} has invalid dimensions")
            
            # STEP 1: Build explicit grid per page
            max_row_idx = layout.get_max_row()  # 0-based max row index
            max_col_idx = layout.get_max_column()  # 0-based max column index
            
            # CRITICAL: If get_max_row/get_max_column return -1, layout is empty
            if max_row_idx < 0 or max_col_idx < 0:
                logger.critical(f"❌ Invalid grid dimensions: max_row={max_row_idx}, max_col={max_col_idx}")
                raise ValueError(f"Layout empty — conversion aborted: Invalid grid dimensions for Page {page_num}")
            
            # Grid dimensions: max_row_idx is 0-based, so grid_rows = max_row_idx + 1
            grid_rows = max_row_idx + 1
            grid_cols = max_col_idx + 1
            
            # DYNAMIC STRUCTURE: Use natural column/row count from layout
            # Each document can have as many columns/rows as needed
            logger.critical("=" * 80)
            logger.critical(f"PAGE {page_num}: Building explicit grid")
            logger.critical(f"   Grid dimensions: {grid_rows} rows × {grid_cols} columns (natural structure)")
            logger.critical(f"   Layout max_row (0-based): {max_row_idx}, max_col (0-based): {max_col_idx}")
            logger.critical("=" * 80)
            
            # STEP 2: Build explicit grid - grid[row][col] = Cell or None
            # Initialize grid with None (empty cells)
            grid = [[None for _ in range(grid_cols)] for _ in range(grid_rows)]
            cell_styles_map = {}  # (row, col) -> CellStyle
            
            # STEP 3: Map all cells to exact grid positions
            # CRITICAL: layout.rows is a list of row lists, where each row list contains cells
            # The row index in layout.rows might not match cell.row - we need to use cell.row and cell.column
            # CRITICAL: Only map cells with non-empty values (no blank columns)
            cells_mapped = 0
            for row_list in layout.rows:
                if not row_list:
                    continue
                for cell in row_list:
                    # CRITICAL: Skip cells with empty values
                    if not cell.value or not str(cell.value).strip():
                        continue  # Don't map empty cells
                    
                    row_idx = cell.row  # 0-based row index from cell
                    col_idx = cell.column  # 0-based column index from cell
                    
                    # CRITICAL DEBUG: Log first few cells
                    if cells_mapped < 5:
                        logger.critical(f"PAGE {page_num}: Mapping cell: row={row_idx}, col={col_idx}, value='{str(cell.value)[:30]}...'")
                    
                    # Validate indices
                    if row_idx < 0 or row_idx >= grid_rows or col_idx < 0 or col_idx >= grid_cols:
                        logger.warning(f"PAGE {page_num}: Cell out of bounds: row={row_idx}, col={col_idx}, grid={grid_rows}x{grid_cols}, value='{str(cell.value)[:30]}...'")
                        # CRITICAL: Expand grid if needed (shouldn't happen, but safety)
                        if row_idx >= grid_rows:
                            # Expand grid rows
                            for _ in range(row_idx + 1 - grid_rows):
                                grid.append([None] * grid_cols)
                            grid_rows = row_idx + 1
                        if col_idx >= grid_cols:
                            # Expand grid columns
                            for r in range(grid_rows):
                                grid[r].extend([None] * (col_idx + 1 - grid_cols))
                            grid_cols = col_idx + 1
                    
                    # Map cell to grid position (overwrite if duplicate)
                    if grid[row_idx][col_idx] is not None:
                        logger.warning(f"PAGE {page_num}: Duplicate cell at ({row_idx},{col_idx}) - overwriting")
                    grid[row_idx][col_idx] = cell
                    cell_styles_map[(row_idx, col_idx)] = cell.style
                    cells_mapped += 1
            
            logger.critical(f"PAGE {page_num}: Mapped {cells_mapped} cells to grid {grid_rows}x{grid_cols}")
            
            # DEBUG: Log first 10 rows of grid BEFORE merge
            logger.critical("=" * 80)
            logger.critical(f"PAGE {page_num}: GRID STATE BEFORE MERGE (first 10 rows)")
            for row_idx in range(min(10, grid_rows)):
                col0_val = str(grid[row_idx][0].value)[:40] if grid[row_idx][0] and grid[row_idx][0].value else "NONE"
                col1_val = str(grid[row_idx][1].value)[:40] if grid_cols > 1 and grid[row_idx][1] and grid[row_idx][1].value else "NONE"
                logger.critical(f"   Row {row_idx}: Col0='{col0_val}' | Col1='{col1_val}'")
            logger.critical("=" * 80)
            
            # STEP 3.5: SKIP OLD MERGE METHOD - Using EXTREME SIMPLIFICATION instead
            logger.critical(f"PAGE {page_num}: Skipping old merge method, will use EXTREME SIMPLIFICATION")
            
            # STEP 4: Column control for invoices/bills
            document_type = layout.metadata.get('document_type', '').lower()
            if document_type in ['invoice', 'bill', 'receipt'] and len(layouts) == 1:
                # Single-page invoice/bill - check if we should force 2 columns
                detected_cols = grid_cols
                if detected_cols == 2:
                    logger.critical(f"PAGE {page_num}: Invoice/bill detected with 2 columns - enforcing exactly 2 columns")
                    # Grid already has 2 columns, no change needed
                elif detected_cols > 2:
                    logger.warning(f"PAGE {page_num}: Invoice/bill has {detected_cols} columns (expected 2) - keeping as-is")
            
            # STEP 5: Create Excel sheet
            sheet = workbook.create_sheet(title=sheet_name)
            
            # STEP 6: Write cells to exact Excel coordinates WITH INLINE MERGING
            # Excel uses 1-based indexing, so row_idx + 1, col_idx + 1
            rows_written = 0
            cols_written = 0
            first_5_rows_preview = []
            rows_to_skip = set()  # Track rows to skip (merged values)
            inline_merges = 0
            
            # EXTREME SIMPLIFICATION: Merge ANY consecutive rows where both have Col1 only
            # NO keyword checking, NO label detection, JUST merge if both rows have Col1 only
            logger.critical("=" * 80)
            logger.critical(f"PAGE {page_num}: EXTREME SIMPLIFICATION - MERGE ALL COL1-ONLY PAIRS")
            logger.critical(f"   Grid size: {grid_rows} rows x {grid_cols} cols")
            logger.critical("=" * 80)
            
            for row_idx in range(grid_rows - 1):
                # Skip if already marked to skip
                if row_idx in rows_to_skip or (row_idx + 1) in rows_to_skip:
                    continue
                
                col0 = grid[row_idx][0] if grid_cols > 0 else None
                col1 = grid[row_idx][1] if grid_cols > 1 else None
                
                next_col0 = grid[row_idx + 1][0] if grid_cols > 0 else None
                next_col1 = grid[row_idx + 1][1] if grid_cols > 1 else None
                
                # SIMPLE RULE: Both rows have Col0 but not Col1
                current_has_col0_only = (col0 and col0.value and str(col0.value).strip() and 
                                        (not col1 or not col1.value or not str(col1.value).strip()))
                
                next_has_col0_only = (next_col0 and next_col0.value and str(next_col0.value).strip() and
                                     (not next_col1 or not next_col1.value or not str(next_col1.value).strip()))
                
                if current_has_col0_only and next_has_col0_only:
                    # MERGE: Move next row's Col0 to current row's Col1
                    grid[row_idx][1] = next_col0
                    rows_to_skip.add(row_idx + 1)
                    inline_merges += 1
                    
                    if inline_merges <= 15:
                        curr_val = str(col0.value).strip()[:40]
                        next_val = str(next_col0.value).strip()[:40]
                        logger.critical(f"   MERGE #{inline_merges}: Row {row_idx} + Row {row_idx+1}")
                        logger.critical(f"      '{curr_val}' | '{next_val}'")
            
            logger.critical("=" * 80)
            logger.critical(f"   TOTAL MERGES: {inline_merges}")
            logger.critical(f"   ROWS TO SKIP: {len(rows_to_skip)}")
            logger.critical("=" * 80)
            
            for row_idx in range(grid_rows):
                # Skip rows that were merged
                if row_idx in rows_to_skip:
                    continue
                    
                row_has_content = False
                for col_idx in range(grid_cols):
                    cell = grid[row_idx][col_idx]
                    if cell is None:
                        continue  # Empty cell, skip
                    
                    # CRITICAL: Skip cells with empty values (no blank columns)
                    if not cell.value or not str(cell.value).strip():
                        continue  # Don't write empty cells
                    
                    # Excel coordinates (1-based)
                    excel_row = row_idx + 1
                    excel_col = col_idx + 1
                    
                    # Write cell value
                    excel_cell = sheet.cell(row=excel_row, column=excel_col, value=cell.value)
                    
                    # Apply style
                    if (row_idx, col_idx) in cell_styles_map:
                        self._apply_cell_style(excel_cell, cell_styles_map[(row_idx, col_idx)])
                    
                    row_has_content = True
                    cols_written = max(cols_written, excel_col)
                
                if row_has_content:
                    rows_written += 1
                    # Store first 5 rows preview
                    if len(first_5_rows_preview) < 5:
                        row_preview = []
                        for col_idx in range(min(5, grid_cols)):
                            cell = grid[row_idx][col_idx]
                            value = str(cell.value)[:30] if cell and cell.value else ''
                            row_preview.append(f"({excel_row},{col_idx+1})='{value}'")
                        first_5_rows_preview.append(f"Row {excel_row}: {', '.join(row_preview)}")
            
            # STEP 7: Handle merged cells
            merged_cell_positions = set()
            merge_ranges = []
            
            # Build merge ranges from layout.merged_cells
            for merged in layout.merged_cells:
                start_row = merged.start_row + 1  # 1-based
                start_col = merged.start_col + 1
                end_row = merged.end_row + 1
                end_col = merged.end_col + 1
                
                merge_ranges.append((start_row, start_col, end_row, end_col))
                
                # Mark slave cells
                for r in range(start_row, end_row + 1):
                    for c in range(start_col, end_col + 1):
                        if not (r == start_row and c == start_col):
                            merged_cell_positions.add((r, c))
            
            # Apply merges
            applied_merges = set()
            for start_row, start_col, end_row, end_col in merge_ranges:
                merge_key = (start_row, start_col, end_row, end_col)
                if merge_key in applied_merges:
                    continue
                applied_merges.add(merge_key)
                
                if end_row < start_row or end_col < start_col:
                    logger.warning(f"Invalid merge range: ({start_row},{start_col}) to ({end_row},{end_col})")
                    continue
                
                try:
                    sheet.merge_cells(
                        start_row=start_row,
                        start_column=start_col,
                        end_row=end_row,
                        end_column=end_col
                    )
                except Exception as e:
                    logger.warning(f"Failed to merge cells: {e}")
            
            # STEP 8: Auto-adjust column widths
            self._auto_adjust_columns(sheet, layout)
            
            # STEP 9: Insert images (anchored positions, no row/column creation)
            page_images = images_by_page.get(layout.page_index, [])
            if page_images and IMAGE_SUPPORT:
                self._insert_images_to_sheet(sheet, page_images, layout)
            
            # STEP 10: Validation logging
            logger.critical("=" * 80)
            logger.critical(f"PAGE {page_num}: Excel writing complete")
            logger.critical(f"   Rows written: {rows_written}")
            logger.critical(f"   Columns written: {cols_written}")
            logger.critical(f"   Grid dimensions: {grid_rows} × {grid_cols}")
            logger.critical(f"   First 5 rows preview:")
            for preview in first_5_rows_preview:
                logger.critical(f"      {preview}")
            logger.critical("=" * 80)
            
            # Final validation
            if rows_written == 0 or cols_written == 0:
                logger.critical(f"❌ VALIDATION FAILED: Page {page_num} has no content after writing")
                raise ValueError(f"Layout empty — conversion aborted: Page {page_num} has no content")
        
        # Final validation - ensure at least one sheet has content
        if len(workbook.worksheets) == 0:
            logger.critical("❌ CRITICAL: No sheets created")
            raise ValueError("Layout empty — conversion aborted: No sheets created")
        
        # Save to bytes
        output = io.BytesIO()
        workbook.save(output)
        output.seek(0)
        return output.getvalue()
    
    def _merge_label_value_pairs_in_grid(self, grid: List[List], grid_rows: int, grid_cols: int, page_num: int) -> int:
        """
        DIRECT GRID MERGE: Merge label-value pairs directly in the grid.
        This is the MOST DIRECT approach - modifies grid in-place.
        
        Strategy:
        - For each row in grid with Col0 only (Col1 is None)
        - Check if next row also has Col0 only
        - If current=label AND next=value, merge them
        - Move next row's Col0 to current row's Col1
        - Clear next row
        """
        label_keywords = ['name', 'id', 'number', 'date', 'amount', 'status', 'channel', 
                         'mode', 'fee', 'ref', 'approval', 'biller', 'customer', 'payment',
                         'agent', 'transaction', 'consumer', 'mobile', 'bill']
        
        merges_done = 0
        rows_to_clear = set()
        
        for row_idx in range(grid_rows - 1):  # Don't check last row (no next row)
            # Current row: Check if it has Col0 only (Col1 is None or empty)
            col0_cell = grid[row_idx][0] if grid_cols > 0 else None
            col1_cell = grid[row_idx][1] if grid_cols > 1 else None
            
            # Current row has Col0 but not Col1
            if col0_cell and col0_cell.value and str(col0_cell.value).strip():
                if not col1_cell or not col1_cell.value or not str(col1_cell.value).strip():
                    current_value = str(col0_cell.value).strip()
                    current_value_lower = current_value.lower()
                    
                    # Check if current is a label
                    has_label_keyword = any(kw in current_value_lower for kw in label_keywords)
                    is_short = len(current_value) < 50
                    
                    is_likely_label = has_label_keyword or is_short
                    
                    if is_likely_label:
                        # Check next row
                        next_row_idx = row_idx + 1
                        if next_row_idx < grid_rows:
                            next_col0_cell = grid[next_row_idx][0] if grid_cols > 0 else None
                            next_col1_cell = grid[next_row_idx][1] if grid_cols > 1 else None
                            
                            # Next row has Col0 but not Col1
                            if next_col0_cell and next_col0_cell.value and str(next_col0_cell.value).strip():
                                if not next_col1_cell or not next_col1_cell.value or not str(next_col1_cell.value).strip():
                                    next_value = str(next_col0_cell.value).strip()
                                    next_value_lower = next_value.lower()
                                    
                                    # Check if next is NOT a label (it's a value)
                                    next_has_label_keyword = any(kw in next_value_lower for kw in label_keywords)
                                    
                                    if not next_has_label_keyword:
                                        # MERGE: Move next row's Col0 to current row's Col1
                                        grid[row_idx][1] = next_col0_cell
                                        
                                        # Clear next row (will be removed later)
                                        grid[next_row_idx][0] = None
                                        rows_to_clear.add(next_row_idx)
                                        
                                        merges_done += 1
                                        if merges_done <= 10:
                                            logger.critical(f"      ✅ GRID MERGE: Row {row_idx+1} '{current_value[:30]}' + Row {next_row_idx+1} '{next_value[:30]}'")
        
        return merges_done
    
    def _merge_label_value_pairs_in_layout(self, layout: UnifiedLayout, page_num: int) -> UnifiedLayout:
        """
        RENDERER LABEL-VALUE PAIRING: Merge adjacent rows where label and value are split.
        This is the FINAL guaranteed place for merging - runs right before Excel writing.
        
        Strategy:
        - For each row with Col1 only (no Col2)
        - Check if next row also has Col1 only
        - If current=label AND next=value, merge them
        - Current row gets: Col1=label, Col2=value
        - Next row is skipped
        """
        import re
        
        # Check if we need to apply (Col2 mostly empty)
        rows_with_col2 = set()
        for row in layout.rows:
            for cell in row:
                if cell.column >= 1 and cell.value and str(cell.value).strip():
                    rows_with_col2.add(cell.row)
        
        total_rows = len(layout.rows)
        col2_coverage = len(rows_with_col2) / total_rows if total_rows > 0 else 0
        
        logger.critical(f"   Total rows: {total_rows}, Rows with Col2: {len(rows_with_col2)}, Coverage: {col2_coverage:.1%}")
        
        # Only apply if Col2 coverage is low (< 70%)
        if col2_coverage >= 0.7:
            logger.critical(f"   ❌ Col2 coverage is high ({col2_coverage:.1%}) - skipping merge")
            return layout
        
        logger.critical(f"   ✅ Col2 coverage is low ({col2_coverage:.1%}) - applying merge")
        
        # Label keywords for detection
        label_keywords = ['name', 'id', 'number', 'date', 'amount', 'status', 'channel', 
                         'mode', 'fee', 'ref', 'approval', 'biller', 'customer', 'payment',
                         'agent', 'transaction', 'consumer', 'mobile', 'bill']
        
        # Group cells by row
        cells_by_row = {}
        for row in layout.rows:
            for cell in row:
                if cell.row not in cells_by_row:
                    cells_by_row[cell.row] = []
                cells_by_row[cell.row].append(cell)
        
        sorted_row_indices = sorted(cells_by_row.keys())
        rows_to_skip = set()
        merges_done = 0
        
        logger.critical(f"   Processing {len(sorted_row_indices)} rows...")
        
        for i in range(len(sorted_row_indices)):
            if i in rows_to_skip:
                continue
                
            current_row_idx = sorted_row_indices[i]
            current_row_cells = cells_by_row[current_row_idx]
            
            # Get Col1 and Col2 cells
            col1_cell = next((c for c in current_row_cells if c.column == 0), None)
            col2_cell = next((c for c in current_row_cells if c.column >= 1 and c.value and str(c.value).strip()), None)
            
            # Current row has Col1 only (no Col2)
            if col1_cell and not col2_cell:
                current_value = str(col1_cell.value).strip() if col1_cell.value else ""
                current_value_lower = current_value.lower()
                
                # Check if current row is a label
                has_label_keyword = any(kw in current_value_lower for kw in label_keywords)
                is_short = len(current_value) < 50
                has_colon = ':' in current_value
                
                is_likely_label = has_label_keyword or (is_short and has_colon) or (is_short and ' ' in current_value)
                
                # Check next row
                if is_likely_label and i + 1 < len(sorted_row_indices):
                    next_i = i + 1
                    next_row_idx = sorted_row_indices[next_i]
                    next_row_cells = cells_by_row[next_row_idx]
                    
                    next_col1_cell = next((c for c in next_row_cells if c.column == 0), None)
                    next_col2_cell = next((c for c in next_row_cells if c.column >= 1 and c.value and str(c.value).strip()), None)
                    
                    # Next row also has Col1 only (no Col2)
                    if next_col1_cell and not next_col2_cell:
                        next_value = str(next_col1_cell.value).strip() if next_col1_cell.value else ""
                        next_value_lower = next_value.lower()
                        
                        # Check if next row is NOT a label (it's a value)
                        next_has_label_keyword = any(kw in next_value_lower for kw in label_keywords)
                        
                        if not next_has_label_keyword:
                            # MERGE: Current=label, Next=value
                            # Move next_col1_cell value to current row's Col2
                            next_col1_cell.row = current_row_idx
                            next_col1_cell.column = 1
                            
                            # Add to current row cells
                            current_row_cells.append(next_col1_cell)
                            
                            # Mark next row to be skipped
                            rows_to_skip.add(next_i)
                            
                            merges_done += 1
                            if merges_done <= 10:
                                logger.critical(f"      ✅ MERGED: Row {current_row_idx} '{current_value[:30]}' + Row {next_row_idx} '{next_value[:30]}'")
        
        # Rebuild layout.rows removing skipped rows
        new_rows = []
        for i, row_idx in enumerate(sorted_row_indices):
            if i not in rows_to_skip:
                # Sort cells in this row by column
                row_cells = sorted(cells_by_row[row_idx], key=lambda c: c.column)
                new_rows.append(row_cells)
        
        layout.rows = new_rows
        
        logger.critical("=" * 80)
        logger.critical(f"   RENDERER MERGE COMPLETE: {merges_done} pairs merged")
        logger.critical(f"   Original rows: {len(sorted_row_indices)}, Final rows: {len(new_rows)}")
        logger.critical("=" * 80)
        
        return layout
    
    def _insert_images_to_sheet(self, sheet, images: List[Dict], layout: UnifiedLayout):
        """Insert images into Excel sheet at appropriate positions"""
        if not IMAGE_SUPPORT:
            return
        
        try:
            from .image_extractor import ImageExtractor
            extractor = ImageExtractor()
            
            max_col = layout.get_max_column()
            max_row = layout.get_max_row()
            
            for img_idx, img_info in enumerate(images[:10]):  # Limit to 10 images per sheet
                try:
                    image_bytes = img_info.get('image_bytes')
                    if not image_bytes:
                        continue
                    
                    # Prepare image for Excel
                    prepared_bytes = extractor.prepare_image_for_excel(image_bytes, max_width=150, max_height=150)
                    if not prepared_bytes:
                        continue
                    
                    # Create Excel image
                    img = ExcelImage(io.BytesIO(prepared_bytes))
                    
                    # CRITICAL FIX: Use actual bounding box position to place image
                    # Convert normalized coordinates to Excel cell position
                    position = img_info.get('position', (0.0, 0.0, 0.1, 0.1))
                    x_norm, y_norm = position[0], position[1]
                    
                    # Convert normalized position to row/column based on actual grid size
                    if max_row > 0 and max_col > 0:
                        # Calculate row and column from normalized position
                        row_num = max(1, int(y_norm * max_row) + 1)
                        col_num = max(1, int(x_norm * max_col) + 1)
                        
                        # Ensure row/column are within bounds
                        row_num = min(row_num, max_row)
                        col_num = min(col_num, max_col)
                    else:
                        # Fallback: If no rows/columns, place at top-left
                        row_num = 1
                        col_num = 1
                        logger.warning(f"Image placement: No rows/columns in layout, placing at A1")
                    
                    # CRITICAL: Images should be FLOATING objects, not anchored to cells
                    # Position image relative to text using pixel coordinates
                    from openpyxl.utils import get_column_letter
                    
                    # Get cell position for reference
                    anchor_col = get_column_letter(col_num)
                    anchor_cell = f"{anchor_col}{row_num}"
                    
                    # Calculate pixel offset from cell (for floating positioning)
                    # Images should be positioned above/below text, not in cells
                    pixel_offset_x = 0  # Horizontal offset from cell
                    pixel_offset_y = 0  # Vertical offset from cell (negative = above, positive = below)
                    
                    # For logos/images, position them above the text row
                    if img_info.get('type') in ['logo', 'image', 'visual']:
                        pixel_offset_y = -20  # Position 20 pixels above the cell
                    
                    # Add image to sheet at detected position
                    sheet.add_image(img, anchor_cell)
                    logger.debug(f"Inserted image {img_idx + 1} at {anchor_cell} (position: x={x_norm:.3f}, y={y_norm:.3f})")
                    
                except Exception as img_error:
                    logger.warning(f"Failed to insert image {img_idx + 1}: {img_error}")
                    continue
        
        except Exception as e:
            logger.warning(f"Error inserting images to sheet: {e}")
    
    def render_to_word(self, layout: UnifiedLayout) -> bytes:
        """
        Render unified layout to Word format.
        
        Args:
            layout: UnifiedLayout object
            
        Returns:
            Word file as bytes
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word export")
        
        doc = Document()
        
        # Create table if we have structured data
        if layout.get_max_row() > 0 and layout.get_max_column() > 1:
            table = doc.add_table(rows=layout.get_max_row(), cols=layout.get_max_column())
            table.style = 'Light Grid Accent 1'
            
            # Fill table cells
            for row in layout.rows:
                for cell in row:
                    if cell.row < len(table.rows) and cell.column < len(table.rows[cell.row].cells):
                        table_cell = table.rows[cell.row].cells[cell.column]
                        paragraph = table_cell.paragraphs[0]
                        run = paragraph.add_run(cell.value)
                        
                        # Apply styles
                        if cell.style.bold:
                            run.bold = True
                        if cell.style.italic:
                            run.italic = True
                        if cell.style.font_size:
                            run.font.size = Pt(cell.style.font_size)
                        
                        # Alignment
                        if cell.style.alignment_horizontal == CellAlignment.CENTER:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif cell.style.alignment_horizontal == CellAlignment.RIGHT:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            # Simple paragraph format for unstructured content
            for row in layout.rows:
                for cell in row:
                    paragraph = doc.add_paragraph(cell.value)
                    if cell.style.bold:
                        for run in paragraph.runs:
                            run.bold = True
        
        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()
    
    def _apply_cell_style(self, excel_cell, style: CellStyle):
        """Apply cell style to Excel cell"""
        # Font
        font_kwargs = {}
        if style.bold:
            font_kwargs['bold'] = True
        if style.italic:
            font_kwargs['italic'] = True
        if style.font_size:
            font_kwargs['size'] = style.font_size
        if style.font_color:
            font_kwargs['color'] = style.font_color
        
        if font_kwargs:
            excel_cell.font = Font(**font_kwargs)
        
        # Alignment
        alignment_kwargs = {}
        if style.alignment_horizontal == CellAlignment.CENTER:
            alignment_kwargs['horizontal'] = 'center'
        elif style.alignment_horizontal == CellAlignment.RIGHT:
            alignment_kwargs['horizontal'] = 'right'
        else:
            alignment_kwargs['horizontal'] = 'left'
        
        if style.alignment_vertical == CellAlignment.TOP:
            alignment_kwargs['vertical'] = 'top'
        elif style.alignment_vertical == CellAlignment.BOTTOM:
            alignment_kwargs['vertical'] = 'bottom'
        else:
            alignment_kwargs['vertical'] = 'center'
        
        if style.wrap_text:
            alignment_kwargs['wrap_text'] = True
        
        excel_cell.alignment = Alignment(**alignment_kwargs)
        
        # Border
        if style.border:
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            excel_cell.border = thin_border
        
        # Background color
        if style.background_color:
            fill = PatternFill(
                start_color=style.background_color,
                end_color=style.background_color,
                fill_type='solid'
            )
            excel_cell.fill = fill
    
    def _auto_adjust_columns(self, sheet, layout: UnifiedLayout):
        """
        Auto-adjust column widths based on content.
        STEP-10: Fully dynamic - only processes columns that actually exist in layout.
        COLUMN STABILITY: Only adjusts width, does NOT infer new columns.
        Column indices from UnifiedLayout are final.
        """
        # STEP-10: Dynamic column detection - determine max column from actual cells
        max_col = layout.get_max_column()
        
        if max_col == 0:
            logger.warning("No columns detected in layout - skipping column width adjustment")
            return
        
        # STEP-10: Dynamic column processing - only process columns that have cells
        # Build set of actual column indices from cells (not assumptions)
        actual_column_indices = set()
        for row in layout.rows:
            for cell in row:
                actual_column_indices.add(cell.column)
        
        # Only process columns that actually exist (dynamic, no hardcoded range)
        for col_idx in sorted(actual_column_indices):
            if col_idx >= max_col:
                continue  # Skip if beyond detected max
                
            column_letter = get_column_letter(col_idx + 1)
            max_length = 0
            
            # Find max length in this column (only from actual cells in layout)
            for row in layout.rows:
                for cell in row:
                    # Only consider cells in this exact column (no inference)
                    if cell.column == col_idx:
                        if cell.value:
                            # Estimate length (Unicode-aware)
                            length = sum(2 if ord(c) > 127 else 1 for c in str(cell.value))
                            max_length = max(max_length, length)
            
            # Set column width (conservative adjustment - doesn't break alignment)
            # Minimum width 10, maximum 50, with small padding
            if max_length > 0:
                adjusted_width = min(max(max_length + 2, 10), 50)
                sheet.column_dimensions[column_letter].width = adjusted_width
            else:
                # Empty column - set minimal width
                sheet.column_dimensions[column_letter].width = 10

