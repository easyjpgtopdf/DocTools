"""
PDF to Word conversion module.
Handles PDF text detection and conversion using LibreOffice or Document AI.
"""
import subprocess
import os
import logging
from typing import Tuple, Optional
from pathlib import Path
from pdfminer.high_level import extract_text as pdf_extract_text
from pdfminer.layout import LAParams
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.docai_client import process_pdf_to_layout, ParsedDocument

logger = logging.getLogger(__name__)


def pdf_has_text(pdf_path: str, max_pages_to_check: int = 3) -> bool:
    """
    Check if PDF contains extractable text.
    
    Args:
        pdf_path: Path to PDF file
        max_pages_to_check: Maximum number of pages to check
        
    Returns:
        True if PDF has extractable text, False otherwise
    """
    try:
        logger.info(f"Checking if PDF has text: {pdf_path}")
        
        # Extract text from first few pages
        laparams = LAParams()
        text = pdf_extract_text(
            pdf_path,
            maxpages=max_pages_to_check,
            laparams=laparams
        )
        
        # Clean and check if meaningful text exists
        cleaned_text = text.strip().replace('\n', '').replace(' ', '')
        
        has_text = len(cleaned_text) > 50  # At least 50 characters
        
        logger.info(f"PDF has text: {has_text} (found {len(cleaned_text)} characters)")
        return has_text
        
    except Exception as e:
        logger.error(f"Error checking PDF text: {e}")
        # If we can't determine, assume it doesn't have text (safer to use DocAI)
        return False


def convert_pdf_to_docx_with_libreoffice(pdf_path: str, output_dir: str) -> str:
    """
    Convert PDF to DOCX using LibreOffice headless (soffice).
    
    This function uses LibreOffice's command-line interface to convert
    PDF files with extractable text to DOCX format. LibreOffice is ideal
    for PDFs that already contain text layers.
    
    Example subprocess command executed:
        soffice --headless --convert-to docx --outdir /tmp input.pdf
    
    Implementation details:
        - Uses subprocess.run() to execute soffice command
        - Runs in headless mode (no GUI)
        - Converts PDF to DOCX format
        - Saves output to specified directory
        - Handles timeouts (5 minutes) and errors gracefully
    
    Args:
        pdf_path: Path to input PDF file
        output_dir: Directory to save output DOCX
        
    Returns:
        Path to created DOCX file
        
    Raises:
        Exception: If conversion fails or times out
    """
    try:
        logger.info(f"Converting PDF to DOCX with LibreOffice: {pdf_path}")
        
        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)
        
        # Build LibreOffice command
        # --headless: Run without GUI
        # --convert-to docx: Convert to DOCX format
        # --outdir: Output directory
        cmd = [
            "soffice",              # LibreOffice executable
            "--headless",           # No GUI mode
            "--convert-to", "docx", # Target format
            "--outdir", output_dir, # Output directory
            pdf_path                # Input file
        ]
        
        logger.debug(f"Executing command: {' '.join(cmd)}")
        
        # Run LibreOffice with timeout protection
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=300  # 5 minute timeout for large files
        )
        
        # Check if command succeeded
        if result.returncode != 0:
            error_msg = result.stderr or result.stdout or "Unknown error"
            raise Exception(f"LibreOffice conversion failed: {error_msg}")
        
        # LibreOffice creates output with same base name as input
        # e.g., input.pdf -> input.docx
        pdf_name = Path(pdf_path).stem
        output_path = os.path.join(output_dir, f"{pdf_name}.docx")
        
        # Verify output file was created
        if not os.path.exists(output_path):
            raise Exception(f"Output DOCX not found at: {output_path}")
        
        logger.info(f"Conversion successful: {output_path}")
        return output_path
        
    except subprocess.TimeoutExpired:
        logger.error("LibreOffice conversion timed out after 5 minutes")
        raise Exception("Conversion timed out. File may be too large or complex.")
    except FileNotFoundError:
        logger.error("LibreOffice (soffice) not found in PATH")
        raise Exception("LibreOffice is not installed or not in PATH")
    except Exception as e:
        logger.error(f"LibreOffice conversion error: {e}")
        raise


def convert_pdf_to_docx_with_docai(
    pdf_path: str,
    output_path: str,
    parsed_doc: ParsedDocument
) -> str:
    """
    Convert PDF to DOCX using Document AI parsed content.
    
    This function takes structured content extracted by Document AI (from scanned PDFs)
    and creates a Word document using python-docx. It preserves:
    - Page structure (page breaks)
    - Text paragraphs
    - Headings (based on font size detection)
    - Tables (if detected by Document AI)
    
    Example Document AI loop over paragraphs:
        for page_num, page_blocks in enumerate(parsed_doc.pages, start=1):
            for block in page_blocks:
                if block.is_heading:
                    doc.add_heading(block.text, level=2)
                else:
                    doc.add_paragraph(block.text)
    
    Implementation details:
        - Creates a new Document() using python-docx
        - Iterates over pages and text blocks from Document AI
        - Preserves headings based on font size heuristics
        - Extracts and formats tables if present
        - Saves the final DOCX document
    
    Args:
        pdf_path: Path to input PDF file (for reference only)
        output_path: Path to save output DOCX file
        parsed_doc: ParsedDocument object from Document AI containing structured content
        
    Returns:
        Path to created DOCX file
        
    Raises:
        Exception: If DOCX creation fails
    """
    try:
        logger.info(f"Creating DOCX from Document AI results: {output_path}")
        
        # Create new Word document
        doc = Document()
        
        # Process each page from Document AI
        for page_num, page_blocks in enumerate(parsed_doc.pages, start=1):
            # Add page break for pages after the first
            if page_num > 1:
                doc.add_page_break()
            
            # Add page header for clarity
            page_heading = doc.add_heading(f"Page {page_num}", level=2)
            page_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Process text blocks on this page
            for block in page_blocks:
                if block.is_heading:
                    # Determine heading level based on font size
                    # Larger fonts (>=18pt) -> Level 1, smaller -> Level 2
                    heading_level = 1 if block.font_size and block.font_size >= 18 else 2
                    doc.add_heading(block.text, level=heading_level)
                else:
                    # Add as regular paragraph
                    paragraph = doc.add_paragraph(block.text)
                    
                    # Apply font size if available (clamp between 8-72pt)
                    if block.font_size:
                        for run in paragraph.runs:
                            run.font.size = Pt(max(8, min(block.font_size, 72)))
            
            # Add tables for this page if any
            for table_data in parsed_doc.tables:
                if table_data.page_number == page_num and table_data.rows:
                    # Create table with appropriate dimensions
                    num_rows = len(table_data.rows)
                    num_cols = len(table_data.rows[0]) if table_data.rows else 0
                    
                    if num_rows > 0 and num_cols > 0:
                        table = doc.add_table(rows=num_rows, cols=num_cols)
                        
                        # Populate table cells
                        for row_idx, row_data in enumerate(table_data.rows):
                            if row_idx < len(table.rows):
                                for col_idx, cell_text in enumerate(row_data):
                                    if col_idx < len(table.rows[row_idx].cells):
                                        table.rows[row_idx].cells[col_idx].text = cell_text
        
        # Save the document
        doc.save(output_path)
        logger.info(f"DOCX created successfully: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"Error creating DOCX from Document AI: {e}")
        raise


def get_page_count(pdf_path: str) -> int:
    """
    Get page count from PDF file.
    
    Args:
        pdf_path: Path to PDF file
        
    Returns:
        Number of pages in the PDF
    """
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(pdf_path)
        return len(reader.pages)
    except ImportError:
        # Fallback to pdfminer if PyPDF2 is not available
        try:
            from pdfminer.high_level import extract_pages
            pages = list(extract_pages(pdf_path))
            return len(pages)
        except Exception as e:
            logger.warning(f"Could not determine page count: {e}")
            return 0
    except Exception as e:
        logger.warning(f"Error getting page count: {e}")
        return 0


def convert_pdf_to_docx(
    pdf_path: str,
    output_dir: str,
    use_docai: bool = False,
    parsed_doc: Optional[ParsedDocument] = None
) -> Tuple[str, bool]:
    """
    Main conversion function. Chooses appropriate method.
    
    Args:
        pdf_path: Path to input PDF
        output_dir: Directory for output DOCX
        use_docai: Force use of Document AI
        parsed_doc: Pre-parsed Document AI result (optional)
        
    Returns:
        Tuple of (output_path, used_docai)
    """
    os.makedirs(output_dir, exist_ok=True)
    pdf_name = Path(pdf_path).stem
    
    if use_docai or parsed_doc:
        # Use Document AI path
        if not parsed_doc:
            raise ValueError("ParsedDocument required when use_docai=True")
        
        output_path = os.path.join(output_dir, f"{pdf_name}.docx")
        convert_pdf_to_docx_with_docai(pdf_path, output_path, parsed_doc)
        return output_path, True
    else:
        # Use LibreOffice path
        output_path = convert_pdf_to_docx_with_libreoffice(pdf_path, output_dir)
        return output_path, False

